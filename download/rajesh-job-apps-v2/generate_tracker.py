#!/usr/bin/env python3
"""
Generate Rajesh Kantubhukta Job Application Tracker - comprehensive DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
pf = style.paragraph_format
pf.space_after = Pt(4)
pf.line_spacing = 1.3

# Colors
NAVY = RGBColor(0x0f, 0x4c, 0x75)
BLUE = RGBColor(0x32, 0x82, 0xb8)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xff, 0xff, 0xff)

# Header
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run('RAJESH KANTUBHUKTA')
r.font.size = Pt(22)
r.font.color.rgb = NAVY
r.font.bold = True

h2 = doc.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = h2.add_run('Job Application Tracker — May 2026')
r2.font.size = Pt(14)
r2.font.color.rgb = BLUE
r2.font.bold = True

h3 = doc.add_paragraph()
h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = h3.add_run('Director of Revenue Cycle Operations | 16+ Years Healthcare RCM + AI')
r3.font.size = Pt(11)
r3.font.color.rgb = GRAY

# Contact info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = contact.add_run('sriraajj@gmail.com | +91 9000111537 | rajeshkantubhukta.github.io | LinkedIn/AgentWorksLab')
r4.font.size = Pt(9)
r4.font.color.rgb = GRAY

doc.add_paragraph()  # spacer

# ==========================================
# SECTION 1: APPLICATION STATUS SUMMARY
# ==========================================
sec1 = doc.add_paragraph()
r = sec1.add_run('APPLICATION STATUS SUMMARY')
r.font.size = Pt(14)
r.font.color.rgb = NAVY
r.font.bold = True

# Summary stats
p = doc.add_paragraph()
r = p.add_run('Total Applications Submitted: 6  |  Manual Applications Needed: 6  |  Total Opportunities: 12')
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = DARK

# Submitted applications table
table1 = doc.add_table(rows=7, cols=5)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Company', 'Role', 'Package', 'Status', 'How Applied']
data = [
    ['HealthRecon Connect', 'VP RCM Operations', '$150K-$180K Remote', 'SUBMITTED (3 roles)', 'Browser Automation'],
    ['Genesis Orthopedics', 'Director of Revenue Cycle', 'India Remote', 'SUBMITTED', 'Workable (Browser)'],
    ['Freenome', 'VP Revenue Cycle Management', 'Remote + 25% travel', 'SUBMITTED', 'Greenhouse (Browser)'],
    ['NIVA Health', 'Revenue Cycle Manager', '$75K+ Remote', 'SUBMITTED', 'Workable (Browser)'],
    ['R1 RCM', 'Director RCM Operations', '$110K-$148K / L30-50L', 'PARTIAL (resume uploaded)', 'Workday (needs manual)'],
    ['Plutus Health', 'Director RCM', 'L35L-L55L Hybrid Hyd', 'NEEDS MANUAL', 'Website form'],
]

for i, h_text in enumerate(headers):
    cell = table1.cell(0, i)
    cell.text = h_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0f4c75"/>')
    cell._tc.get_or_add_tcPr().append(shading)

for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table1.cell(row_idx + 1, col_idx)
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
        # Color status column
        if col_idx == 3 and 'SUBMITTED' in val:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0x0a, 0x80, 0x43)
                run.font.bold = True

doc.add_paragraph()

# ==========================================
# SECTION 2: MANUAL APPLICATIONS NEEDED
# ==========================================
sec2 = doc.add_paragraph()
r = sec2.add_run('MANUAL APPLICATIONS NEEDED (Apply Links + Instructions)')
r.font.size = Pt(14)
r.font.color.rgb = NAVY
r.font.bold = True

p = doc.add_paragraph()
r = p.add_run('These companies have complex Workday or custom application forms that require manual submission. Your tailored resume and cover letter PDFs are included in the download folder.')
r.font.size = Pt(10)
r.font.color.rgb = GRAY

manual_apps = [
    {
        'company': 'CorroHealth',
        'role': 'VP RCM Operations',
        'location': 'Noida, India',
        'why': 'AI analytics + RCM. Your exact VP-level role at an AI-powered RCM company. Technology-driven transformation.',
        'url': 'https://corrohealth.wd1.myworkdayjobs.com/en-US/CorroHealthIndia/job/VP-RCM-Operations_JR104933',
        'resume': 'Resume_CorroHealth.pdf',
        'cover': 'Cover_CorroHealth.pdf',
        'notes': 'Workday application - need to create account first. Use sriraajj@gmail.com. Job ID: JR104933.'
    },
    {
        'company': 'U.S. Urology Partners',
        'role': 'VP Revenue Cycle Management',
        'location': 'US Remote (H1B Sponsor!)',
        'why': 'H1B visa sponsor confirmed. $190K-$250K package. Specialty care RCM. This is a high-value target.',
        'url': 'https://usuro.wd1.myworkdayjobs.com/en-US/USUrology_Careers/job/Vice-President--Revenue-Cycle-Management_R2750',
        'resume': 'Resume_USUrology.pdf',
        'cover': 'Cover_USUrology.pdf',
        'notes': 'Workday application. H1B sponsor - explicitly mention willingness to relocate with visa support. Job ID: R2750.'
    },
    {
        'company': 'Huron Consulting',
        'role': 'VP Managed Services Revenue Cycle Operations',
        'location': 'US Remote',
        'why': 'Consulting mindset role. 10+ years leadership required. Healthcare BPO expertise valued. $200K+ range.',
        'url': 'https://huron.wd1.myworkdayjobs.com/en-US/huroncareers/job/Vice-President--Managed-Services-Revenue-Cycle-Operations_JR-0013577',
        'resume': 'Resume_Huron.pdf',
        'cover': 'Cover_Huron.pdf',
        'notes': 'Workday application. Job ID: JR-0013577. Emphasize consulting + managed services delivery experience.'
    },
    {
        'company': 'Plutus Health',
        'role': 'Director RCM',
        'location': 'Hybrid Hyderabad',
        'why': 'Inc. 5000 fastest growing. Your exact title. AI in RCM focus. SOC + HIPAA certified.',
        'url': 'https://www.plutushealthinc.com/careers/director-rcm',
        'resume': 'Resume_Genesis.pdf (use as base)',
        'cover': 'Cover_Genesis.pdf (use as base)',
        'notes': 'Website form. Also has: Director QA-RCM, AM/Manager AR/RCM. Apply to all matching roles.'
    },
    {
        'company': 'RethinkFirst',
        'role': 'Director of RCM Operations',
        'location': 'Remote (offshore teams)',
        'why': 'Managing offshore teams is your core strength. Billing/collections/QA. $77K-$139K.',
        'url': 'https://builtin.com/job/director-rcm-operations/9244598',
        'resume': 'Resume_RethinkFirst.pdf',
        'cover': 'Cover_RethinkFirst.pdf',
        'notes': 'Apply via BuiltIn or ZipRecruiter. Emphasize offshore team management (120+ FTEs).'
    },
    {
        'company': 'Optum India',
        'role': 'Associate Director RCM Operations',
        'location': 'Hybrid Hyderabad',
        'why': 'UHG brand. Massive scale. Career launcher. L35L-L55L.',
        'url': 'https://www.optum.in/about/careers.html',
        'resume': 'Resume_Genesis.pdf (use as base)',
        'cover': 'Cover_Genesis.pdf (use as base)',
        'notes': 'Search for Director/Associate Director RCM roles on their careers page.'
    },
]

for app in manual_apps:
    p = doc.add_paragraph()
    r = p.add_run(f"{app['company']}  —  {app['role']}")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = NAVY
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Location: {app['location']}")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GRAY
    
    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Why You: {app['why']}")
    r3.font.size = Pt(10)
    
    p4 = doc.add_paragraph()
    r4a = p4.add_run('Apply URL: ')
    r4a.font.bold = True
    r4a.font.size = Pt(9.5)
    r4b = p4.add_run(app['url'])
    r4b.font.size = Pt(9.5)
    r4b.font.color.rgb = BLUE
    
    p5 = doc.add_paragraph()
    r5a = p5.add_run('Resume: ')
    r5a.font.bold = True
    r5a.font.size = Pt(9.5)
    r5b = p5.add_run(app['resume'])
    r5b.font.size = Pt(9.5)
    r5c = p5.add_run('   |   Cover Letter: ')
    r5c.font.bold = True
    r5c.font.size = Pt(9.5)
    r5d = p5.add_run(app['cover'])
    r5d.font.size = Pt(9.5)
    
    p6 = doc.add_paragraph()
    r6a = p6.add_run('Notes: ')
    r6a.font.bold = True
    r6a.font.size = Pt(9.5)
    r6b = p6.add_run(app['notes'])
    r6b.font.size = Pt(9.5)
    r6b.font.color.rgb = GRAY
    
    # Add separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(6)
    r_sep = sep.add_run('─' * 80)
    r_sep.font.size = Pt(6)
    r_sep.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

doc.add_paragraph()

# ==========================================
# SECTION 3: JOB PORTALS FOR RECRUITER VISIBILITY
# ==========================================
sec3 = doc.add_paragraph()
r = sec3.add_run('JOB PORTALS — Get Recruiters Coming to You')
r.font.size = Pt(14)
r.font.color.rgb = NAVY
r.font.bold = True

p = doc.add_paragraph()
r = p.add_run('Upload your resume to these portals so recruiters find you. This is the highest-ROI action for getting calls without applying one-by-one.')
r.font.size = Pt(10)
r.font.color.rgb = GRAY

portals = [
    {
        'name': 'Naukri.com (PRIORITY #1)',
        'url': 'https://www.naukri.com/mnjuser/profile',
        'action': 'Create profile as "Director - Revenue Cycle Management" with keywords: US Healthcare RCM, Denial Management, Revenue Cycle Operations, AI in Healthcare, Remote. Set job alerts for Director RCM, VP RCM. 32,000+ RCM jobs listed.',
        'impact': 'HIGH - India\'s #1 job site. Recruiters actively search here. 470+ Director RCM jobs currently open.'
    },
    {
        'name': 'LinkedIn (PRIORITY #2)',
        'url': 'https://www.linkedin.com/jobs/',
        'action': 'Update headline to: "Director of Revenue Cycle Operations | 16+ Years US Healthcare RCM | AI Product Builder | Open to Remote". Turn on "Open to Work" with Remote + India locations. Set job alerts for VP/Director RCM Remote.',
        'impact': 'HIGH - 1,295+ Director Revenue Cycle Remote jobs. Recruiters proactively search LinkedIn profiles.'
    },
    {
        'name': 'Indeed India',
        'url': 'https://in.indeed.com/',
        'action': 'Upload resume. Set job alerts for "Director Revenue Cycle Remote India". 1,295+ Director Revenue Cycle remote jobs listed.',
        'impact': 'MEDIUM - Large volume but less recruiter-driven than Naukri/LinkedIn.'
    },
    {
        'name': 'ZipRecruiter',
        'url': 'https://www.ziprecruiter.com/',
        'action': 'Create profile. 1,000+ Remote Director RCM jobs ($87K-$182K). AI matches you to jobs.',
        'impact': 'MEDIUM - Good for US remote roles. AI-driven matching means recruiters find you.'
    },
    {
        'name': 'Glassdoor India',
        'url': 'https://www.glassdoor.co.in/',
        'action': 'Upload resume. 381 Revenue Cycle Director remote jobs. Also check company reviews before applying.',
        'impact': 'MEDIUM - Good for salary research and company intelligence.'
    },
]

for portal in portals:
    p = doc.add_paragraph()
    r = p.add_run(portal['name'])
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = NAVY
    
    p2 = doc.add_paragraph()
    r2a = p2.add_run('URL: ')
    r2a.font.bold = True
    r2b = p2.add_run(portal['url'])
    r2b.font.color.rgb = BLUE
    r2b.font.size = Pt(9.5)
    
    p3 = doc.add_paragraph()
    r3a = p3.add_run('Action: ')
    r3a.font.bold = True
    r3b = p3.add_run(portal['action'])
    r3b.font.size = Pt(10)
    
    p4 = doc.add_paragraph()
    r4a = p4.add_run('Impact: ')
    r4a.font.bold = True
    r4a.font.color.rgb = RGBColor(0x0a, 0x80, 0x43)
    r4b = p4.add_run(portal['impact'])
    r4b.font.size = Pt(10)
    r4b.font.color.rgb = RGBColor(0x0a, 0x80, 0x43)

doc.add_paragraph()

# ==========================================
# SECTION 4: COMPANIES TO MONITOR
# ==========================================
sec4 = doc.add_paragraph()
r = sec4.add_run('COMPANIES TO MONITOR (Check Weekly for New Openings)')
r.font.size = Pt(14)
r.font.color.rgb = NAVY
r.font.bold = True

monitor_companies = [
    ('AKASA', 'AI for healthcare operations. Remote-friendly. Check akasa.com/careers', 'https://akasa.com/careers'),
    ('FinThrive', 'Has India careers page. Revenue management focus. Check finthrive.com/careers-india', 'https://finthrive.com/careers-india'),
    ('Infinx Healthcare', 'AI-driven RCM. Patient access + billing. Check infinix.com/careers', 'https://www.infinx.com/careers'),
    ('GeBBS Healthcare', 'Major India RCM employer. Check careers.gebbs.com', 'https://careers.gebbs.com/gebbs/jobslist/India'),
    ('Omega Healthcare', 'Largest India RCM company. Check omegahms.com/careers-india', 'https://www.omegahms.com/careers-india'),
    ('Access Healthcare', 'End-to-end RCM services. Check accesshealthcare.com', 'https://www.accesshealthcare.com'),
    ('Commure', 'Healthcare technology + AI. Check commure.com/careers', 'https://www.commure.com/careers'),
    ('WNS Global', 'Healthcare BPO + RCM. Check wns.com/careers', 'https://www.wns.com/careers'),
    ('Genpact Healthcare', 'Healthcare analytics + RCM. Check genpact.com/careers', 'https://www.genpact.com/careers'),
    ('EXL Healthcare', 'Analytics + operations. Check exlservice.com/careers', 'https://www.exlservice.com/careers'),
]

table2 = doc.add_table(rows=len(monitor_companies) + 1, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h_text in enumerate(['Company', 'Why Monitor', 'Careers URL']):
    cell = table2.cell(0, i)
    cell.text = h_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0f4c75"/>')
    cell._tc.get_or_add_tcPr().append(shading)

for row_idx, (name, why, url) in enumerate(monitor_companies):
    table2.cell(row_idx + 1, 0).text = name
    table2.cell(row_idx + 1, 1).text = why
    table2.cell(row_idx + 1, 2).text = url
    for col in range(3):
        for paragraph in table2.cell(row_idx + 1, col).paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ==========================================
# SECTION 5: YOUR UNIQUE PITCH
# ==========================================
sec5 = doc.add_paragraph()
r = sec5.add_run('YOUR UNIQUE PITCH (Use in Every Application)')
r.font.size = Pt(14)
r.font.color.rgb = NAVY
r.font.bold = True

pitch = doc.add_paragraph()
r = pitch.add_run('"I\'m an RCM operator who has built a production AI platform (Denial Doctor) with 6 specialist agents, NCCI validation, corrected code generation, and verified-citation appeal letters. No other candidate has this combination of 16+ years operational depth + AI product engineering. I manage 120+ FTEs, $500M+ in annual revenue, have recovered $47M+ in denied revenue, and reduced Days in A/R from 65 to 42. Published author of \'The $50K Leak\' on Amazon Kindle. My story: rajeshkantubhukta.github.io"')
r.font.size = Pt(10.5)
r.font.italic = True
r.font.color.rgb = DARK

doc.add_paragraph()

# Key metrics box
p = doc.add_paragraph()
r = p.add_run('KEY METRICS TO HIGHLIGHT IN EVERY APPLICATION:')
r.font.bold = True
r.font.color.rgb = NAVY

metrics = [
    '16+ years end-to-end US Healthcare RCM',
    '$500M+ annual revenue managed',
    '120+ FTEs across Patient Access, Coding, Billing, Collections',
    '98% client retention rate',
    '$47M+ denied revenue recovered',
    'Days in A/R reduced from 65 to 42',
    '75% client portfolio growth',
    'Denial Doctor: 320% ROI, $12M annualized savings, 6 AI agents',
    'Published Author: "The $50K Leak" (Amazon Kindle, 2026)',
    'AI Platform: HIPAA-compliant, FHIR R4, anti-hallucination guardrails',
]

for m in metrics:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(m)
    r.font.size = Pt(10)

# Save
output_path = '/home/z/my-project/download/Rajesh_Kantubhukta_Job_Application_Tracker_May2026.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
